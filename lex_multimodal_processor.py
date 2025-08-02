#!/usr/bin/env python3
"""
LEX Multimodal Processor - Handles various file types and routes to appropriate models
"""
import os
import base64
import mimetypes
import aiofiles
from typing import Dict, List, Any, Optional, Tuple
from pathlib import Path
import PyPDF2
import docx
import pandas as pd
from PIL import Image
import io
import json

# Import enhanced PDF processor
try:
    from enhanced_pdf_processor import enhanced_pdf_processor
    HAS_ENHANCED_PDF = True
except ImportError:
    HAS_ENHANCED_PDF = False
    print("⚠️ Enhanced PDF processor not available, using basic PDF processing")

class MultimodalProcessor:
    def __init__(self):
        # Models that support vision/multimodal
        self.vision_capable_models = {
            "llava:13b": {
                "supports": ["image"],
                "quality": 0.85,
                "speed": 0.4
            },
            "bakllava:latest": {
                "supports": ["image"],
                "quality": 0.8,
                "speed": 0.5
            },
            "llava:7b": {
                "supports": ["image"],
                "quality": 0.75,
                "speed": 0.7
            }
        }
        
        # File type handlers
        self.file_handlers = {
            "image": self.process_image,
            "pdf": self.process_pdf,
            "document": self.process_document,
            "spreadsheet": self.process_spreadsheet,
            "video": self.process_video,
            "text": self.process_text
        }
        
        # MIME type mapping
        self.mime_mapping = {
            # Images
            "image/jpeg": "image",
            "image/png": "image",
            "image/gif": "image",
            "image/webp": "image",
            "image/bmp": "image",
            
            # Documents
            "application/pdf": "pdf",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "document",
            "application/msword": "document",
            "application/vnd.oasis.opendocument.text": "document",
            
            # Spreadsheets
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": "spreadsheet",
            "application/vnd.ms-excel": "spreadsheet",
            "text/csv": "spreadsheet",
            
            # Videos
            "video/mp4": "video",
            "video/mpeg": "video",
            "video/quicktime": "video",
            "video/x-msvideo": "video",
            "video/webm": "video",
            
            # Text
            "text/plain": "text",
            "text/markdown": "text",
            "application/json": "text",
            "application/xml": "text"
        }
        
        print("Multimodal Processor initialized")
    
    def get_file_type(self, mime_type: str) -> str:
        """Determine file type from MIME type"""
        return self.mime_mapping.get(mime_type, "text")
    
    async def process_file(self, file_path: str, mime_type: str = None) -> Dict[str, Any]:
        """Process any file and extract content"""
        if not mime_type:
            mime_type, _ = mimetypes.guess_type(file_path)
        
        file_type = self.get_file_type(mime_type)
        handler = self.file_handlers.get(file_type, self.process_text)
        
        try:
            result = await handler(file_path)
            result["file_type"] = file_type
            result["mime_type"] = mime_type
            return result
        except Exception as e:
            return {
                "error": f"Failed to process {file_type}: {str(e)}",
                "file_type": file_type,
                "mime_type": mime_type
            }
    
    async def process_image(self, file_path: str) -> Dict[str, Any]:
        """Process image files"""
        try:
            # Use enhanced processor for OCR if available
            ocr_text = ""
            if HAS_ENHANCED_PDF:
                ocr_result = await enhanced_pdf_processor.process_image(
                    file_path,
                    extract_images=False,
                    extract_tables=False,
                    ocr_mode="auto"
                )
                if ocr_result.get('success'):
                    ocr_text = ocr_result.get('text', '')
            
            # Open and analyze image
            with Image.open(file_path) as img:
                # Get image info
                width, height = img.size
                format = img.format
                mode = img.mode
                
                # Convert to base64 for vision models
                buffered = io.BytesIO()
                img.save(buffered, format=format or 'PNG')
                img_base64 = base64.b64encode(buffered.getvalue()).decode('utf-8')
                
                # Generate thumbnail for preview
                img.thumbnail((300, 300))
                thumb_buffered = io.BytesIO()
                img.save(thumb_buffered, format='PNG')
                thumb_base64 = base64.b64encode(thumb_buffered.getvalue()).decode('utf-8')
            
            return {
                "success": True,
                "content_type": "image",
                "data": {
                    "base64": img_base64,
                    "thumbnail": thumb_base64,
                    "width": width,
                    "height": height,
                    "format": format,
                    "mode": mode,
                    "ocr_text": ocr_text  # Add OCR text if available
                },
                "metadata": {
                    "dimensions": f"{width}x{height}",
                    "format": format,
                    "size_kb": os.path.getsize(file_path) / 1024,
                    "has_text": bool(ocr_text)
                },
                "requires_vision_model": True
            }
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    async def process_pdf(self, file_path: str) -> Dict[str, Any]:
        """Process PDF files"""
        try:
            # Use enhanced PDF processor if available
            if HAS_ENHANCED_PDF:
                result = await enhanced_pdf_processor.process_pdf(
                    file_path,
                    extract_images=True,
                    extract_tables=True,
                    ocr_mode="auto"
                )
                
                if result.get('success'):
                    # Format for compatibility with existing code
                    return {
                        "success": True,
                        "content_type": "pdf",
                        "data": {
                            "text": result.get('text', '')[:20000],
                            "pages": result.get('pages', [])[:10],
                            "total_pages": result.get('total_pages', 0),
                            "images": result.get('images', []),
                            "tables": result.get('tables', [])
                        },
                        "metadata": {
                            "ocr_performed": result.get('ocr_performed', False),
                            "pages": result.get('total_pages', 0),
                            "processing_details": result.get('processing_details', {})
                        },
                        "summary": result.get('summary', ''),
                        "requires_vision_model": False
                    }
            
            # Fallback to basic PDF processing
            text_content = []
            metadata = {}
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                num_pages = len(pdf_reader.pages)
                
                # Extract text from each page
                for page_num, page in enumerate(pdf_reader.pages):
                    text = page.extract_text()
                    if text.strip():
                        text_content.append({
                            "page": page_num + 1,
                            "text": text.strip()
                        })
                
                # Get metadata
                if pdf_reader.metadata:
                    metadata = {
                        "title": pdf_reader.metadata.get('/Title', ''),
                        "author": pdf_reader.metadata.get('/Author', ''),
                        "subject": pdf_reader.metadata.get('/Subject', ''),
                        "pages": num_pages
                    }
            
            # Combine text for processing
            full_text = "\n\n".join([f"Page {p['page']}:\n{p['text']}" for p in text_content])
            
            return {
                "success": True,
                "content_type": "pdf",
                "data": {
                    "text": full_text[:20000],  # Increased limit for better context
                    "pages": text_content[:10],  # First 10 pages
                    "total_pages": num_pages
                },
                "metadata": metadata,
                "requires_vision_model": False
            }
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    async def process_document(self, file_path: str) -> Dict[str, Any]:
        """Process Word documents"""
        try:
            doc = docx.Document(file_path)
            
            # Extract text
            full_text = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text.strip())
            
            # Extract tables
            tables = []
            for table in doc.tables[:5]:  # Limit to first 5 tables
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)
            
            return {
                "success": True,
                "content_type": "document",
                "data": {
                    "text": "\n\n".join(full_text)[:10000],
                    "paragraphs": len(full_text),
                    "tables": tables
                },
                "metadata": {
                    "paragraphs": len(full_text),
                    "tables": len(doc.tables)
                },
                "requires_vision_model": False
            }
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    async def process_spreadsheet(self, file_path: str) -> Dict[str, Any]:
        """Process Excel/CSV files"""
        try:
            # Determine file type
            if file_path.endswith('.csv'):
                df = pd.read_csv(file_path)
            else:
                df = pd.read_excel(file_path)
            
            # Get basic info
            rows, cols = df.shape
            
            # Convert to analyzable format
            summary = {
                "columns": df.columns.tolist(),
                "shape": f"{rows} rows x {cols} columns",
                "dtypes": df.dtypes.to_dict(),
                "sample": df.head(10).to_dict(),
                "statistics": df.describe().to_dict() if rows > 0 else {}
            }
            
            # Convert to markdown table for preview
            preview_md = df.head(10).to_markdown()
            
            return {
                "success": True,
                "content_type": "spreadsheet",
                "data": {
                    "preview": preview_md,
                    "summary": summary,
                    "full_data": df.to_dict() if rows < 1000 else None
                },
                "metadata": {
                    "rows": rows,
                    "columns": cols
                },
                "requires_vision_model": False
            }
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    async def process_video(self, file_path: str) -> Dict[str, Any]:
        """Process video files (basic metadata for now)"""
        try:
            # For now, just get basic file info
            # In future, could extract frames or use video models
            file_size = os.path.getsize(file_path)
            
            return {
                "success": True,
                "content_type": "video",
                "data": {
                    "note": "Video processing requires specialized models. Currently extracting metadata only."
                },
                "metadata": {
                    "size_mb": file_size / (1024 * 1024),
                    "path": file_path
                },
                "requires_vision_model": True,
                "requires_special_model": "video"
            }
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    async def process_text(self, file_path: str) -> Dict[str, Any]:
        """Process plain text files"""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
                content = await f.read()
            
            # Limit size for context
            if len(content) > 50000:
                content = content[:50000] + "\n\n[Content truncated...]"
            
            return {
                "success": True,
                "content_type": "text",
                "data": {
                    "text": content
                },
                "metadata": {
                    "length": len(content),
                    "lines": content.count('\n') + 1
                },
                "requires_vision_model": False
            }
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    def select_vision_model(self, file_type: str, available_models: List[str]) -> Optional[str]:
        """Select best vision model for file type"""
        if file_type not in ["image", "video"]:
            return None
        
        # Check which vision models are available
        for model, info in self.vision_capable_models.items():
            if model in available_models and file_type in info["supports"]:
                return model
        
        return None
    
    def prepare_multimodal_prompt(self, text_prompt: str, file_data: Dict[str, Any]) -> str:
        """Prepare prompt that includes file context"""
        file_type = file_data.get("content_type", "unknown")
        
        if file_type == "image":
            return f"[Image provided: {file_data['metadata']['dimensions']}, {file_data['metadata']['format']}]\n\n{text_prompt}"
        
        elif file_type == "pdf":
            pages = file_data['data'].get('total_pages', 0)
            # Get more meaningful text content
            text_content = file_data['data'].get('text', '')
            if len(text_content) > 10000:
                text_content = text_content[:10000] + "..."
            
            # Create a comprehensive prompt for PDF analysis
            return f"""You have been provided with a PDF document ({pages} pages).

PDF Content:
{text_content}

User Request: {text_prompt}

IMPORTANT INSTRUCTIONS:
1. Provide a COMPREHENSIVE and DETAILED summary/analysis
2. Include ALL key information from the document
3. Use specific numbers, dates, names, and facts
4. Organize your response with clear sections
5. Do NOT provide short or incomplete responses
6. Your response should be at least several paragraphs long

Begin your comprehensive analysis:"""
        
        elif file_type == "document":
            return f"[Word Document]\n\nContent:\n{file_data['data']['text'][:2000]}...\n\nUser Query: {text_prompt}"
        
        elif file_type == "spreadsheet":
            return f"[Spreadsheet Data]\n\n{file_data['data']['preview']}\n\nUser Query: {text_prompt}"
        
        elif file_type == "text":
            return f"[Text File]\n\nContent:\n{file_data['data']['text'][:2000]}...\n\nUser Query: {text_prompt}"
        
        else:
            return f"[File: {file_type}]\n\n{text_prompt}"


# Singleton instance
multimodal_processor = MultimodalProcessor()